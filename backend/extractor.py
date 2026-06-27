import io
import re
import zipfile
from pathlib import PurePosixPath
from typing import Iterable, Tuple

import pandas as pd
import pdfplumber
from docx import Document


SUPPORTED_EXTENSIONS = (".pdf", ".xlsx", ".xls", ".csv", ".docx", ".txt")
GENERIC_ZIP_ROOTS = {"", "архив", "archive", "prices", "price", "прайсы", "прайс", "хакатон"}


def detect_file_format(filename: str, file_bytes: bytes | None = None) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith((".xlsx", ".xls")):
        return "xlsx"
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith(".csv"):
        return "csv"
    if lower.endswith(".zip"):
        return "zip"
    return "text"


def _clean_archive_root(value: str) -> str:
    text = re.sub(r"[_-]+", " ", value or "")
    text = re.sub(r"(?i)\b(202\d|20\d\d|год|zip|archive|архив)\b", " ", text)
    return re.sub(r"\s+", " ", text).strip(" ._-–—").lower()


def _clean_display_filename(path: str) -> str:
    """Normalize display path while keeping the price year visible.

    `Клиника 2 прайс 2025 год.PDF` -> `Клиника 2 прайс 2025.PDF`.
    This keeps `2025` for the UI, but prevents partner inference from returning
    `Клиника 2 год` after it removes `прайс` and the numeric year.
    """
    p = PurePosixPath(path)
    stem = p.stem
    suffix = p.suffix
    stem = re.sub(r"(?i)\b(20\d{2})\s*год\b", r"\1", stem)
    stem = re.sub(r"(?i)\bгод\b", " ", stem)
    stem = re.sub(r"\s+", " ", stem).strip(" ._-–—")
    cleaned_name = f"{stem}{suffix}"
    if p.parent == PurePosixPath("."):
        return cleaned_name
    return str(p.parent / cleaned_name)


def _strip_generic_zip_root(zip_filename: str, inner_paths: list[str]) -> dict[str, str]:
    """Remove one common technical root folder from ZIP paths.

    Example: `Хакатон/Клиника 1 2026.pdf` becomes `Клиника 1 2026.pdf`,
    so partner inference uses the real clinic file name, not the archive folder.
    Real clinic folders are kept: `Clinic A/price.xlsx`, `Clinic B/price.xlsx`.
    """
    if not inner_paths:
        return {}

    split_paths = [p.split("/") for p in inner_paths]
    if not all(len(parts) > 1 for parts in split_paths):
        return {p: _clean_display_filename(p) for p in inner_paths}

    first_root = split_paths[0][0]
    if not all(parts[0] == first_root for parts in split_paths):
        return {p: _clean_display_filename(p) for p in inner_paths}

    clean_root = _clean_archive_root(first_root)
    clean_zip = _clean_archive_root(PurePosixPath(zip_filename or "").stem)
    should_strip = clean_root in GENERIC_ZIP_ROOTS or clean_root == clean_zip
    if not should_strip:
        return {p: _clean_display_filename(p) for p in inner_paths}

    return {p: _clean_display_filename(str(PurePosixPath(*p.split("/")[1:]))) for p in inner_paths}


def iter_input_files(filename: str, file_bytes: bytes) -> Iterable[Tuple[str, bytes]]:
    """Yield original file or files inside ZIP.

    For ZIP files we keep meaningful inner relative paths. One generic archive
    root folder is stripped, but real clinic folders are preserved.
    """
    if filename.lower().endswith(".zip"):
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
            entries: list[tuple[str, str]] = []
            for info in archive.infolist():
                if info.is_dir():
                    continue
                inner_path = str(PurePosixPath(info.filename.replace("\\", "/")))
                basename = PurePosixPath(inner_path).name
                if not basename or basename.startswith(".") or "__MACOSX" in inner_path:
                    continue
                if not basename.lower().endswith(SUPPORTED_EXTENSIONS):
                    continue
                entries.append((inner_path, info.filename))

            display_paths = _strip_generic_zip_root(filename, [inner for inner, _ in entries])
            for inner_path, archive_name in entries:
                yield display_paths.get(inner_path, _clean_display_filename(inner_path)), archive.read(archive_name)
    else:
        yield _clean_display_filename(filename), file_bytes


def extract_text_from_pdf(file_bytes: bytes) -> str:
    chunks: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                chunks.append(f"\n--- PDF page {page_index} text ---\n{page_text}")

            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []

            for table_index, table in enumerate(tables, start=1):
                rows = []
                for row in table:
                    rows.append("\t".join("" if cell is None else str(cell).strip() for cell in row))
                if rows:
                    chunks.append(f"\n--- PDF page {page_index} table {table_index} ---\n" + "\n".join(rows))

    return "\n".join(chunks).strip()


def extract_text_from_excel(file_bytes: bytes, filename: str = "") -> str:
    lower = filename.lower()
    if lower.endswith(".csv"):
        df = pd.read_csv(io.BytesIO(file_bytes), header=None, dtype=str)
        return df.fillna("").to_string(index=False, header=False)

    sheets = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, header=None, dtype=str)
    chunks: list[str] = []
    for sheet_name, df in sheets.items():
        df = df.dropna(how="all").fillna("")
        if df.empty:
            continue
        lines = []
        for row in df.astype(str).values.tolist():
            cleaned = [cell.strip() for cell in row]
            if any(cleaned):
                lines.append("\t".join(cleaned))
        if lines:
            chunks.append(f"\n--- XLSX sheet: {sheet_name} ---\n" + "\n".join(lines))
    return "\n".join(chunks).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    chunks: list[str] = []

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if paragraphs:
        chunks.append("\n--- DOCX text ---\n" + "\n".join(paragraphs))

    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append("\t".join(cell.text.strip() for cell in row.cells))
        if rows:
            chunks.append(f"\n--- DOCX table {table_index} ---\n" + "\n".join(rows))

    return "\n".join(chunks).strip()


def extract_text(filename: str, file_bytes: bytes) -> str:
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if filename_lower.endswith((".xlsx", ".xls", ".csv")):
        return extract_text_from_excel(file_bytes, filename)
    if filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("cp1251")
        except UnicodeDecodeError as exc:
            raise ValueError(f"Формат файла {filename} не поддерживается.") from exc
